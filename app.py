from flask import Flask, render_template, jsonify, request,  send_file,  redirect, url_for
import json
import io
import PyPDF2
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from IPython.display import display, Markdown
import google.generativeai as genai
import base64
from io import BytesIO
import io
import os
import PIL.Image
from PIL import Image
import textwrap
from docx import Document
from gtts import gTTS
from bs4 import BeautifulSoup
import requests
import pandas as pd
import matplotlib.pyplot as plt


app = Flask(__name__)

genai.configure(api_key="AIzaSyAbmCYsZsjfCPf-uakFksDglYasW4EsehE")


@app.route('/')
def home():
    return render_template('index.html')



@app.route('/chatbot')
def chatbot():
    return render_template('chatbot.html')

@app.route('/get_response', methods=['POST'])
def get_response():
    data = request.json
    user_message = data['message']
    model = genai.GenerativeModel('gemini-1.5-flash')
    rply = model.generate_content(f"{user_message}  answer in one or 2 lines without any */ symbols")

    return jsonify({'response': rply.text})

@app.route('/web')
def web():
    return render_template('web-summary.html')

def summarize_text(text):
    model = genai.GenerativeModel('gemini-1.5-flash')
    rply = model.generate_content("Summarize this webpage and give me the summary of this webpage: " + text)
    to_markdown(rply.text)
    return rply.text

def to_markdown(text):
    text = text.replace('\u2022', '  *')
    return Markdown(textwrap.indent(text, '> ', predicate=lambda _: True))

def scrape_webpage(url):
    try:
        response = requests.get(url)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'html.parser')
        paragraphs = soup.find_all('p')
        text = ' '.join(paragraph.text for paragraph in paragraphs)
        return text
    except requests.RequestException as e:
        return f"An error occurred while fetching the webpage: {e}"

@app.route('/web_summary', methods=['POST'])
def web_summary():
    url = request.form['url']
    webpage_text = scrape_webpage(url)
    websummary = summarize_text(webpage_text)
    tts = gTTS(text=websummary, lang='en')
    output_file_audio = "static/outputsummary.mp3"
    tts.save(output_file_audio)
    return render_template('web-result.html', websummary=websummary, audio_file=output_file_audio)



@app.route('/codemaster')
def code():
    return render_template('code.html')

@app.route('/generate', methods=['POST'])
def generate_html():
    prompt = request.form['prompt']
    model = genai.GenerativeModel('gemini-1.5-flash')
    rply = model.generate_content("Generate both html and css code in single file for " + prompt + " with colorful CSS background and more attractive CSS, I need only code and no explanation")
    html_content = rply.text

    with open("templates/index1.html", "w") as file:
        file.write(html_content)

    return redirect(url_for('output'))


@app.route('/output')
def output():
    return render_template('index1.html')

def strip_code(response_text):
    response_text = response_text.strip()
    if response_text.startswith('```json') and response_text.endswith('```'):
        return response_text[7:-3].strip()
    return response_text

def generate_certificate_pdf(name, certificate_text, template_path, output_path):
    template_pdf = PyPDF2.PdfReader(template_path)
    template_page = template_pdf.pages[0]

    packet = io.BytesIO()
    can = canvas.Canvas(packet, pagesize=letter)

    can.setFont("Helvetica-Bold", 24)
    name_width = can.stringWidth(name, "Helvetica-Bold", 24)
    x = ((letter[0] - name_width) / 2) + 100
    y = letter[1] - 500
    can.drawString(x, y, name)

    can.setFont("Helvetica", 12)
    text_width = can.stringWidth(certificate_text, "Helvetica", 12)
    x = ((letter[0] - text_width) / 2) + 100
    y -= 30
    can.drawString(x, y, certificate_text)

    can.save()
    packet.seek(0)

    new_pdf = PyPDF2.PdfReader(packet)
    output = PyPDF2.PdfWriter()
    template_page.merge_page(new_pdf.pages[0])
    output.add_page(template_page)

    with open(output_path, 'wb') as f:
        output.write(f)


@app.route('/quiz', methods=['GET', 'POST'])
def quiz():
    topic = ""
    questions = []

    if request.method == 'POST':
        topic = request.form.get('topic')
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(
            f"""Create a 5-question multiple-choice quiz on the topic of {topic}. 
               Provide the questions, four answer options for each question, and the correct answer. 
               Format the output as a JSON array with each question as an object containing: 
               'question', 'options' (list of four strings), and 'answer' (string)."""
        )
        response_text = response.text
        stripped_text = strip_code(response_text)

        try:
            questions = json.loads(stripped_text)
        except json.JSONDecodeError as e:
            print(f"Error parsing response: {e}")
            questions = []

    return render_template('quiz.html', topic=topic, questions=questions)

@app.route('/submit', methods=['POST'])
def submit():
    questions = json.loads(request.form.get('questions'))
    user_answers = request.form.to_dict(flat=True)
    user_answers.pop('questions')

    correct_count = 0
    total_questions = len(questions)

    for i, question in enumerate(questions):
        correct_answer = question['answer']
        user_answer = user_answers.get(f'question-{i}')
        if user_answer == correct_answer:
            correct_count += 1

    return render_template('results.html', total=total_questions, correct=correct_count, incorrect=total_questions - correct_count)

@app.route('/certificate', methods=['GET', 'POST'])
def certificate():
    if request.method == 'POST' :
        name = request.form.get('username')
        certificate_text = f"{name} has successfully completed the quiz."
        template_path = 'certificate.pdf'
        output_path = 'certificate_.pdf'

        generate_certificate_pdf(name, certificate_text, template_path, output_path)
        return send_file(output_path, as_attachment=True)

    return render_template('certificate.html')


@app.route('/speech')
def speechindex():
    return render_template('speech.html')

@app.route('/store_transcription', methods=['POST'])
def store_transcription():
    data = request.json
    transcription = data['transcription']

    # Store transcription in a Word file
    document = Document()
    document.add_paragraph(transcription)
    document.save('transcription.docx')

    return 'Transcription stored successfully'

@app.route('/speech-summary')
def speech():
    with open('transcription.docx', 'rb') as docx_file:
        doc = Document(docx_file)
        script_summary = ""
        for para in doc.paragraphs:
            script_summary += para.text + "<br>"
    
    # Render HTML template with the content
    return render_template('transcript-result.html', script_summary=script_summary)

@app.route('/to-do-list')
def to_do_list():
    return render_template('to-do-list.html')

@app.route('/timer')
def timer():
    return render_template('timer.html')

@app.route('/roadmap')
def code1():
    return render_template('road.html')

@app.route('/road', methods=['POST'])
def generate_html1():
    prompt = request.form['prompt']
    model = genai.GenerativeModel('gemini-1.5-flash')
    rply = model.generate_content("Generate both html and css code in single file for Roadmap on topic " + prompt + " with 30 days study plans and resources ,colorful CSS background and more attractive CSS, I need only code and no explanation")
    html_content = rply.text

    with open("templates/index2.html", "w") as file:
        file.write(html_content)

    return redirect(url_for('output1'))


@app.route('/output1')
def output1():
    return render_template('index2.html')

def plot_line_chart(subjects_df):
    df_mean = subjects_df.mean()
    plt.figure()
    df_mean.plot(kind='line', marker='o')
    plt.title('Trend of Average Marks per Subject')
    plt.xlabel('Subjects')
    plt.ylabel('Average Marks')
    line_chart = io.BytesIO()
    plt.savefig(line_chart, format='png')
    line_chart.seek(0)
    plt.close()
    return line_chart

@app.route('/mark', methods=['GET', 'POST'])
def mark():
    if request.method == 'POST':
        if 'file' not in request.files:
            return redirect(request.url)
        file = request.files['file']
        if file.filename == '':
            return redirect(request.url)
        if file:
            df = pd.read_excel(file)
            subjects_df = df.drop(columns=['Roll No', 'Name'])

            # Generate line chart
            line_chart = plot_line_chart(subjects_df)
            line_chart_url = base64.b64encode(line_chart.getvalue()).decode('utf8')

            return render_template('mark.html', line_chart_url=line_chart_url)

    return render_template('mark.html')


@app.route('/navi')
def navi():
    return render_template('navi.html')

@app.route('/gpt', methods=['GET', 'POST'])
def gpt():
    response_text = ""
    audio=''
    if request.method == 'POST':
        # Get transcribed text from the form
        transcribed_text = request.form.get('transcribed_text')
          
        # Generate response using the transcribed text
        if transcribed_text:
            # Generate response using Generative AI model
            model = genai.GenerativeModel('gemini-pro')
            rply = model.generate_content("explain in 3 lines"+ transcribed_text)
            response_text = rply.text
            print(response_text)
            # Convert response text to speech
            tts = gTTS(text=response_text, lang='en')
            tts.save('response.mp3')
            # Encode the audio file as base64
            with open("response.mp3", "rb") as audio_file:
                 encoded_string = base64.b64encode(audio_file.read()).decode('utf-8')
        else:
            response_text = "No input provided."
            encoded_string = ""
        
        # Return the response to the client
        return render_template('gpt.html', response=response_text, audio=encoded_string)
    else:
        # If it's a GET request, render the form
        return render_template('gpt.html')

if __name__ == '__main__':
    app.run(debug=True)
